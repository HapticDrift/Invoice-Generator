import sqlite3
from docx import Document
from docx.shared import Inches
import pandas as pd
import win32com.client

conn  =  sqlite3 . connect ( 'MyInvoice.db' )
cursor  =  conn.cursor ()
#create the salesman table 
cursor.execute("CREATE TABLE sales(name char(50), mailAddress char(50), product char(30), units n(5), price n(7));")


def createInvoice(name, email, product, unit, price):
    document = Document()
    document.add_picture('icons8-trollface-48.png', width=Inches(1))
    document.add_heading('Invoice', 0)
    p1 = document.add_paragraph('Dear ')
    p1.add_run(name).bold=True
    p1.add_run(',')

    p2 = document.add_paragraph('Please find attached invoice for your recent purchase of ')
    p2.add_run(str(unit)).bold = True
    p2.add_run(' units of ')
    p2.add_run(product).bold=True
    p2.add_run('.')

    [document.add_paragraph('') for _ in range(2)]
    
    table = document.add_table(rows=1, cols=4)
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Product Name'
    hdr_cells[1].text = 'Units'
    hdr_cells[2].text = 'Unit Price'
    hdr_cells[3].text = 'Total Price'
    for i in range(4):
        hdr_cells[i].paragraphs[0].runs[0].font.bold = True
        
    row_cells = table.add_row().cells
    row_cells[0].text = product
    row_cells[1].text = f'{unit:,.2f}'
    row_cells[2].text = f'{price:,.2f}'
    row_cells[3].text = f'{unit * price:,.2f}'
    
    [document.add_paragraph('') for _ in range(10)]

    document.add_paragraph('We appreciate you coming to our aid!')
    document.add_paragraph('Sincerely')
    document.add_paragraph('Lord Denathor')

    document.save(f'{name}.docx')

def docx_to_pdf(src, dst):
    word = win32com.client.Dispatch("Word.Application")
    wdFormatPDF = 17
    doc = word.Documents.Open(src)
    doc.SaveAs(dst, FileFormat=wdFormatPDF)
    doc.Close()
    word.Quit()

def send_email(name, to_addr, attachment):
    outlook = win32com.client.Dispatch("Outlook.Application")
    mail = outlook.CreateItem(0)
    mail.To = to_addr #'amznbotnotification@gmail.com'
    mail.Subject = 'Invoice from Minas Tirith'
    mail.Body = f'Dear {name}, Please find attached invoice'
    mail.Attachments.Add(attachment)
    mail.Send()

title = input("Enter your name: ")
mailAddress = input("Please specify your e-mail address: ")
item = input("Item dealing with:" )
units = int(input("Units Purchased: "))
cost = int(input("Cost per item is: "))

cursor.execute("""
INSERT INTO sales(name, mailAddress, product, units, price)
VALUES (?,?,?,?,?)
""", (title, mailAddress, item, units, cost))
conn.commit ()
print ( 'Data entered successfully.' )
conn . close ()
if (conn):
  conn.close()
  print("\nThe SQLite connection is closed.")
createInvoice(title, mailAddress, item, units, cost)